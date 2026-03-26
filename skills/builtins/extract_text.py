"""
文本提取Skill

从各种文档格式中提取文本内容。
"""

import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import io

from ..base import BaseSkill, SkillResult, SkillMetadata, SkillParameter
from ..decorators import skill, parameter, tag, category
from ..context import SkillContext


@skill(
    name="extract_text",
    version="1.0.0",
    description="从文档中提取纯文本内容",
    author="DocMCP Team",
    category="document_processing",
    tags=["text", "extraction", "document"],
    timeout=120.0
)
@parameter("source", Union[str, Path, bytes], "文档来源（文件路径或内容）", required=True)
@parameter("format", str, "文档格式（auto/pdf/docx/txt/html/md）", default="auto")
@parameter("encoding", str, "文本编码", default="utf-8")
@parameter("preserve_layout", bool, "保留布局信息", default=False)
@parameter("extract_metadata", bool, "提取元数据", default=True)
@tag("core", "extraction")
@category("document_processing")
class ExtractTextSkill(BaseSkill):
    """
    文本提取Skill
    
    支持从多种文档格式中提取文本内容：
    - PDF (.pdf)
    - Word (.docx)
    - 纯文本 (.txt)
    - HTML (.html, .htm)
    - Markdown (.md)
    
    示例:
        skill = ExtractTextSkill()
        result = skill.run(context, source="document.pdf")
        print(result.data["text"])
    """
    
    SUPPORTED_FORMATS = {
        "txt": "plain_text",
        "text": "plain_text",
        "pdf": "pdf",
        "docx": "docx",
        "doc": "docx",
        "html": "html",
        "htm": "html",
        "md": "markdown",
        "markdown": "markdown",
        "json": "json",
        "xml": "xml",
    }
    
    def _on_initialize(self, context: SkillContext) -> SkillResult:
        """初始化Skill"""
        context.log_info("ExtractTextSkill 初始化完成")
        return SkillResult.success_result()
    
    def execute(
        self,
        context: SkillContext,
        source: Union[str, Path, bytes],
        format: str = "auto",
        encoding: str = "utf-8",
        preserve_layout: bool = False,
        extract_metadata: bool = True,
        **kwargs
    ) -> SkillResult:
        """
        执行文本提取
        
        Args:
            context: 执行上下文
            source: 文档来源
            format: 文档格式
            encoding: 文本编码
            preserve_layout: 是否保留布局
            extract_metadata: 是否提取元数据
            
        Returns:
            提取结果
        """
        try:
            # 确定格式
            if format == "auto":
                format = self._detect_format(source)
            
            format = format.lower()
            
            # 提取文本
            if format in ["txt", "text", "plain_text"]:
                result = self._extract_plain_text(source, encoding)
            elif format == "pdf":
                result = self._extract_pdf(source, encoding, preserve_layout)
            elif format == "docx":
                result = self._extract_docx(source)
            elif format in ["html", "htm"]:
                result = self._extract_html(source, encoding)
            elif format in ["md", "markdown"]:
                result = self._extract_markdown(source, encoding)
            elif format == "json":
                result = self._extract_json(source, encoding)
            elif format == "xml":
                result = self._extract_xml(source, encoding)
            else:
                return SkillResult.error_result(f"不支持的格式: {format}")
            
            # 构建结果
            output = {
                "text": result.get("text", ""),
                "format": format,
                "encoding": encoding,
            }
            
            if extract_metadata:
                output["metadata"] = result.get("metadata", {})
            
            if preserve_layout:
                output["layout"] = result.get("layout", {})
            
            # 统计信息
            output["stats"] = {
                "char_count": len(output["text"]),
                "line_count": output["text"].count("\n") + 1,
                "word_count": len(output["text"].split()),
            }
            
            context.log_info(f"成功提取文本，共 {output['stats']['char_count']} 字符")
            
            return SkillResult.success_result(data=output)
            
        except Exception as e:
            context.log_error(f"文本提取失败: {str(e)}")
            return SkillResult.error_result(f"文本提取失败: {str(e)}")
    
    def _detect_format(self, source: Union[str, Path, bytes]) -> str:
        """检测文档格式"""
        if isinstance(source, (str, Path)):
            path = Path(source)
            suffix = path.suffix.lower().lstrip(".")
            return self.SUPPORTED_FORMATS.get(suffix, "txt")
        elif isinstance(source, bytes):
            # 尝试从内容检测
            if source.startswith(b"%PDF"):
                return "pdf"
            elif source.startswith(b"PK"):
                return "docx"
            elif source.startswith(b"<?xml"):
                return "xml"
            elif source.startswith(b"{"):
                return "json"
            elif b"<!DOCTYPE html" in source[:1000].lower() or b"<html" in source[:1000].lower():
                return "html"
            return "txt"
        return "txt"
    
    def _read_source(self, source: Union[str, Path, bytes], encoding: str) -> Union[str, bytes]:
        """读取来源内容"""
        if isinstance(source, (str, Path)):
            path = Path(source)
            if path.exists():
                if self._detect_format(source) in ["pdf", "docx"]:
                    return path.read_bytes()
                return path.read_text(encoding=encoding)
            else:
                # 假设是文本内容
                return source
        return source
    
    def _extract_plain_text(
        self,
        source: Union[str, Path, bytes],
        encoding: str
    ) -> Dict[str, Any]:
        """提取纯文本"""
        content = self._read_source(source, encoding)
        
        if isinstance(content, bytes):
            content = content.decode(encoding)
        
        return {
            "text": content,
            "metadata": {
                "type": "plain_text",
                "encoding": encoding
            }
        }
    
    def _extract_pdf(
        self,
        source: Union[str, Path, bytes],
        encoding: str,
        preserve_layout: bool
    ) -> Dict[str, Any]:
        """提取PDF文本"""
        try:
            # 尝试使用PyPDF2
            from PyPDF2 import PdfReader
            
            content = self._read_source(source, encoding)
            if isinstance(content, str):
                content = content.encode(encoding)
            
            reader = PdfReader(io.BytesIO(content))
            
            text_parts = []
            metadata = {}
            
            # 提取元数据
            if reader.metadata:
                metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "subject": reader.metadata.get("/Subject", ""),
                    "creator": reader.metadata.get("/Creator", ""),
                    "producer": reader.metadata.get("/Producer", ""),
                    "pages": len(reader.pages),
                }
            
            # 提取文本
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if preserve_layout:
                    text_parts.append(f"\n--- Page {i + 1} ---\n")
                text_parts.append(page_text or "")
            
            return {
                "text": "\n".join(text_parts),
                "metadata": metadata,
                "layout": {"pages": len(reader.pages)} if preserve_layout else {}
            }
            
        except ImportError:
            # 降级处理
            return self._extract_plain_text(source, encoding)
    
    def _extract_docx(
        self,
        source: Union[str, Path, bytes]
    ) -> Dict[str, Any]:
        """提取Word文档文本"""
        try:
            from docx import Document
            
            content = self._read_source(source, "utf-8")
            if isinstance(content, str):
                content = content.encode("utf-8")
            
            doc = Document(io.BytesIO(content))
            
            # 提取段落
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 提取表格
            tables_text = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_rows.append(" | ".join(row_text))
                tables_text.append("\n".join(table_rows))
            
            # 合并文本
            all_text = "\n\n".join(paragraphs)
            if tables_text:
                all_text += "\n\n" + "\n\n".join(tables_text)
            
            # 提取元数据
            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
            }
            
            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                })
            
            return {
                "text": all_text,
                "metadata": metadata
            }
            
        except ImportError:
            return self._extract_plain_text(source, "utf-8")
    
    def _extract_html(
        self,
        source: Union[str, Path, bytes],
        encoding: str
    ) -> Dict[str, Any]:
        """提取HTML文本"""
        try:
            from bs4 import BeautifulSoup
            
            content = self._read_source(source, encoding)
            if isinstance(content, bytes):
                content = content.decode(encoding)
            
            soup = BeautifulSoup(content, "html.parser")
            
            # 移除script和style
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 提取文本
            text = soup.get_text(separator="\n")
            
            # 清理空白
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)
            
            # 提取元数据
            metadata = {}
            title = soup.find("title")
            if title:
                metadata["title"] = title.get_text()
            
            meta_tags = soup.find_all("meta")
            for meta in meta_tags:
                name = meta.get("name", "").lower()
                content_val = meta.get("content", "")
                if name and content_val:
                    metadata[name] = content_val
            
            return {
                "text": text,
                "metadata": metadata
            }
            
        except ImportError:
            return self._extract_plain_text(source, encoding)
    
    def _extract_markdown(
        self,
        source: Union[str, Path, bytes],
        encoding: str
    ) -> Dict[str, Any]:
        """提取Markdown文本"""
        try:
            import markdown
            from bs4 import BeautifulSoup
            
            content = self._read_source(source, encoding)
            if isinstance(content, bytes):
                content = content.decode(encoding)
            
            # 转换为HTML然后提取文本
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text(separator="\n")
            
            # 提取标题
            headers = re.findall(r'^#{1,6}\s+(.+)$', content, re.MULTILINE)
            
            # 提取链接
            links = re.findall(r'\[([^\]]+)\]\(([^)]+)\)', content)
            
            return {
                "text": text,
                "metadata": {
                    "headers": headers,
                    "link_count": len(links),
                    "links": links[:10]  # 只保留前10个链接
                }
            }
            
        except ImportError:
            return self._extract_plain_text(source, encoding)
    
    def _extract_json(
        self,
        source: Union[str, Path, bytes],
        encoding: str
    ) -> Dict[str, Any]:
        """提取JSON文本"""
        import json
        
        content = self._read_source(source, encoding)
        if isinstance(content, bytes):
            content = content.decode(encoding)
        
        try:
            data = json.loads(content)
            
            # 格式化输出
            formatted = json.dumps(data, indent=2, ensure_ascii=False)
            
            return {
                "text": formatted,
                "metadata": {
                    "type": type(data).__name__,
                    "keys": list(data.keys()) if isinstance(data, dict) else [],
                    "length": len(data) if hasattr(data, "__len__") else 0
                }
            }
            
        except json.JSONDecodeError as e:
            return {
                "text": content,
                "metadata": {"error": f"JSON解析错误: {str(e)}"}
            }
    
    def _extract_xml(
        self,
        source: Union[str, Path, bytes],
        encoding: str
    ) -> Dict[str, Any]:
        """提取XML文本"""
        try:
            import xml.etree.ElementTree as ET
            
            content = self._read_source(source, encoding)
            if isinstance(content, bytes):
                content = content.decode(encoding)
            
            root = ET.fromstring(content)
            
            # 提取所有文本
            texts = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    texts.append(elem.text.strip())
            
            # 提取标签结构
            def get_structure(elem, level=0):
                result = {"tag": elem.tag, "children": []}
                for child in elem:
                    result["children"].append(get_structure(child, level + 1))
                return result
            
            return {
                "text": "\n".join(texts),
                "metadata": {
                    "root_tag": root.tag,
                    "structure": get_structure(root)
                }
            }
            
        except Exception as e:
            return {
                "text": content if isinstance(content, str) else content.decode(encoding),
                "metadata": {"error": f"XML解析错误: {str(e)}"}
            }
