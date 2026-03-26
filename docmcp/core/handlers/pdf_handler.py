"""
PDF文档处理器模块

支持 .pdf 格式的PDF文档处理，包括文本提取、图片提取和元数据读取。
"""

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, BinaryIO, Dict, Iterator, List, Optional, Tuple, Union

from .base import BaseDocument, BaseDocumentHandler, BatchProcessor
from ..document import DocumentMetadata, DocumentType, ExtractedContent
from ..utils import TempFileManager, detect_file_type, chunk_file_reader

# 配置日志
logger = logging.getLogger(__name__)

# 尝试导入依赖库
try:
    import PyPDF2
    from PyPDF2 import PdfReader, PdfWriter
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning("PyPDF2 未安装，PDF功能受限")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber 未安装，PDF文本提取功能受限")

try:
    from pdf2image import convert_from_path, convert_from_bytes
    from PIL import Image
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    logger.warning("pdf2image 未安装，PDF图片转换功能受限")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("Pillow 未安装，图片处理功能受限")


class PDFDocument(BaseDocument):
    """
    PDF文档对象
    
    支持 .pdf 格式的PDF文档处理。
    """
    
    # 大文件阈值 (50MB for PDF)
    LARGE_FILE_THRESHOLD = 50 * 1024 * 1024
    
    def __init__(
        self,
        file_path: Optional[Union[str, Path]] = None,
        file_stream: Optional[BinaryIO] = None,
        document_type: Optional[DocumentType] = None
    ):
        """
        初始化PDF文档对象
        
        Args:
            file_path: 文档文件路径
            file_stream: 文档文件流
            document_type: 文档类型
        """
        super().__init__(file_path, file_stream, document_type)
        self._pdf_reader: Optional[Any] = None
        self._pdfplumber_doc: Optional[Any] = None
        self._page_count: int = 0
        self._temp_manager = TempFileManager()
    
    def load(self, **kwargs) -> "PDFDocument":
        """
        加载PDF文档
        
        Args:
            strict: 是否严格模式（PyPDF2参数）
            
        Returns:
            self: 支持链式调用
        """
        if self._is_loaded:
            return self
        
        strict = kwargs.get('strict', False)
        
        try:
            # 使用PyPDF2加载
            if PYPDF2_AVAILABLE:
                self._load_with_pypdf2(strict=strict)
            
            # 使用pdfplumber加载（更好的文本提取）
            if PDFPLUMBER_AVAILABLE:
                self._load_with_pdfplumber()
            
            self._is_loaded = True
            self._update_metadata_from_file()
            
        except Exception as e:
            logger.error(f"加载PDF文档失败: {e}")
            raise ValueError(f"无法加载PDF文档: {e}")
        
        return self
    
    def _load_with_pypdf2(self, strict: bool = False) -> None:
        """使用PyPDF2加载PDF"""
        if self._file_stream:
            self._file_stream.seek(0)
            self._pdf_reader = PdfReader(self._file_stream, strict=strict)
        elif self._file_path:
            self._pdf_reader = PdfReader(str(self._file_path), strict=strict)
        else:
            raise ValueError("没有可用的文件源")
        
        self._page_count = len(self._pdf_reader.pages)
    
    def _load_with_pdfplumber(self) -> None:
        """使用pdfplumber加载PDF"""
        if self._file_stream:
            self._file_stream.seek(0)
            self._pdfplumber_doc = pdfplumber.open(self._file_stream)
        elif self._file_path:
            self._pdfplumber_doc = pdfplumber.open(self._file_path)
        else:
            raise ValueError("没有可用的文件源")
    
    def get_page_count(self) -> int:
        """
        获取页面数量
        
        Returns:
            页面数量
        """
        if not self._is_loaded:
            self.load()
        return self._page_count
    
    def extract_text(self, **kwargs) -> str:
        """
        提取纯文本内容
        
        Args:
            page_numbers: 指定页面列表（可选，从1开始）
            start_page: 起始页面（从1开始）
            end_page: 结束页面（包含，从1开始）
            preserve_layout: 是否保留布局
            
        Returns:
            文档的纯文本内容
        """
        if not self._is_loaded:
            self.load()
        
        page_numbers = kwargs.get('page_numbers')
        start_page = kwargs.get('start_page', 1)
        end_page = kwargs.get('end_page', self._page_count)
        preserve_layout = kwargs.get('preserve_layout', False)
        
        # 确定要处理的页面
        if page_numbers:
            pages_to_process = [p - 1 for p in page_numbers if 1 <= p <= self._page_count]
        else:
            pages_to_process = range(start_page - 1, min(end_page, self._page_count))
        
        texts = []
        
        # 优先使用pdfplumber提取（效果更好）
        if PDFPLUMBER_AVAILABLE and self._pdfplumber_doc:
            for page_idx in pages_to_process:
                page = self._pdfplumber_doc.pages[page_idx]
                text = page.extract_text()
                if text:
                    texts.append(text)
        
        # 回退到PyPDF2
        elif PYPDF2_AVAILABLE and self._pdf_reader:
            for page_idx in pages_to_process:
                page = self._pdf_reader.pages[page_idx]
                text = page.extract_text()
                if text:
                    texts.append(text)
        
        return '\n\n'.join(texts)
    
    def extract_text_by_pages(self, **kwargs) -> Dict[int, str]:
        """
        按页面提取文本
        
        Returns:
            页面号到文本的字典
        """
        if not self._is_loaded:
            self.load()
        
        result = {}
        
        if PDFPLUMBER_AVAILABLE and self._pdfplumber_doc:
            for page_idx, page in enumerate(self._pdfplumber_doc.pages):
                text = page.extract_text()
                if text:
                    result[page_idx + 1] = text
        
        elif PYPDF2_AVAILABLE and self._pdf_reader:
            for page_idx, page in enumerate(self._pdf_reader.pages):
                text = page.extract_text()
                if text:
                    result[page_idx + 1] = text
        
        return result
    
    def extract_content(self, **kwargs) -> ExtractedContent:
        """
        提取完整内容
        
        Args:
            extract_tables: 是否提取表格
            extract_images: 是否提取图片
            
        Returns:
            包含页面内容的完整内容
        """
        if not self._is_loaded:
            self.load()
        
        extract_tables = kwargs.get('extract_tables', True)
        extract_images = kwargs.get('extract_images', False)
        
        content = ExtractedContent()
        
        if PDFPLUMBER_AVAILABLE and self._pdfplumber_doc:
            for page_idx, page in enumerate(self._pdfplumber_doc.pages):
                page_data = {
                    'index': page_idx,
                    'page_number': page_idx + 1,
                    'text': '',
                    'tables': [],
                    'images': [],
                    'width': page.width,
                    'height': page.height
                }
                
                # 提取文本
                text = page.extract_text()
                if text:
                    page_data['text'] = text
                    content.paragraphs.append(text)
                
                # 提取表格
                if extract_tables:
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            page_data['tables'].append(table)
                            content.tables.append(table)
                
                # 提取图片（需要额外处理）
                if extract_images:
                    # pdfplumber不直接支持图片提取，使用其他方法
                    pass
                
                content.pages.append(page_data)
        
        elif PYPDF2_AVAILABLE and self._pdf_reader:
            for page_idx, page in enumerate(self._pdf_reader.pages):
                page_data = {
                    'index': page_idx,
                    'page_number': page_idx + 1,
                    'text': '',
                    'tables': [],
                    'images': []
                }
                
                text = page.extract_text()
                if text:
                    page_data['text'] = text
                    content.paragraphs.append(text)
                
                content.pages.append(page_data)
        
        content.text = '\n\n'.join(content.paragraphs)
        
        return content
    
    def extract_tables(self, **kwargs) -> List[List[List[str]]]:
        """
        提取所有表格
        
        Args:
            page_numbers: 指定页面列表（可选）
            
        Returns:
            表格数据列表
        """
        if not self._is_loaded:
            self.load()
        
        if not PDFPLUMBER_AVAILABLE:
            logger.warning("需要安装 pdfplumber 库来提取表格")
            return []
        
        page_numbers = kwargs.get('page_numbers')
        tables = []
        
        pages_to_process = (page_numbers if page_numbers 
                          else range(1, self._page_count + 1))
        
        for page_num in pages_to_process:
            if 1 <= page_num <= self._page_count:
                page = self._pdfplumber_doc.pages[page_num - 1]
                page_tables = page.extract_tables()
                tables.extend(page_tables)
        
        return tables
    
    def extract_images(
        self,
        output_dir: Optional[Union[str, Path]] = None,
        **kwargs
    ) -> List[Path]:
        """
        提取PDF中的图片
        
        Args:
            output_dir: 输出目录
            min_width: 最小图片宽度
            min_height: 最小图片高度
            
        Returns:
            提取的图片路径列表
        """
        if not self._is_loaded:
            self.load()
        
        min_width = kwargs.get('min_width', 100)
        min_height = kwargs.get('min_height', 100)
        
        if not output_dir:
            output_dir = self._temp_manager.create_temp_dir()
        else:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
        
        image_paths = []
        
        # 使用pdf2image转换页面为图片
        if PDF2IMAGE_AVAILABLE:
            try:
                if self._file_path:
                    images = convert_from_path(
                        str(self._file_path),
                        dpi=kwargs.get('dpi', 200)
                    )
                elif self._file_stream:
                    self._file_stream.seek(0)
                    images = convert_from_bytes(
                        self._file_stream.read(),
                        dpi=kwargs.get('dpi', 200)
                    )
                else:
                    return image_paths
                
                for idx, image in enumerate(images):
                    image_path = output_dir / f"page_{idx + 1}.png"
                    image.save(image_path, 'PNG')
                    image_paths.append(image_path)
                    
            except Exception as e:
                logger.error(f"pdf2image转换失败: {e}")
        
        return image_paths
    
    def extract_metadata(self, **kwargs) -> DocumentMetadata:
        """
        提取文档元数据
        
        Returns:
            文档元数据对象
        """
        if not self._is_loaded:
            self.load()
        
        metadata = DocumentMetadata()
        
        # 从文件系统获取基本元数据
        self._update_metadata_from_file()
        metadata.file_size = self._metadata.file_size
        metadata.file_path = self._metadata.file_path
        metadata.file_name = self._metadata.file_name
        metadata.file_extension = self._metadata.file_extension
        
        # 从PDF提取元数据
        if PYPDF2_AVAILABLE and self._pdf_reader:
            try:
                pdf_metadata = self._pdf_reader.metadata
                
                if pdf_metadata:
                    metadata.title = pdf_metadata.get('/Title')
                    metadata.author = pdf_metadata.get('/Author')
                    metadata.subject = pdf_metadata.get('/Subject')
                    metadata.keywords = pdf_metadata.get('/Keywords')
                    metadata.creator = pdf_metadata.get('/Creator')
                    metadata.producer = pdf_metadata.get('/Producer')
                    
                    # 日期处理
                    creation_date = pdf_metadata.get('/CreationDate')
                    if creation_date:
                        metadata.created = self._parse_pdf_date(creation_date)
                    
                    mod_date = pdf_metadata.get('/ModDate')
                    if mod_date:
                        metadata.modified = self._parse_pdf_date(mod_date)
                
                # 页面信息
                metadata.custom_properties['page_count'] = self._page_count
                
                # 检查是否加密
                if self._pdf_reader.is_encrypted:
                    metadata.custom_properties['encrypted'] = True
                    
            except Exception as e:
                logger.debug(f"提取PDF元数据失败: {e}")
        
        return metadata
    
    def _parse_pdf_date(self, date_str: str) -> Optional[datetime]:
        """
        解析PDF日期字符串
        
        PDF日期格式: D:YYYYMMDDHHmmSSOHH'mm'
        
        Args:
            date_str: PDF日期字符串
            
        Returns:
            datetime对象或None
        """
        try:
            # 移除前缀 'D:'
            if date_str.startswith('D:'):
                date_str = date_str[2:]
            
            # 提取基本日期部分 (YYYYMMDDHHmmSS)
            if len(date_str) >= 14:
                year = int(date_str[0:4])
                month = int(date_str[4:6])
                day = int(date_str[6:8])
                hour = int(date_str[8:10])
                minute = int(date_str[10:12])
                second = int(date_str[12:14])
                
                return datetime(year, month, day, hour, minute, second)
                
        except Exception as e:
            logger.debug(f"解析PDF日期失败: {e}")
        
        return None
    
    def save(self, output_path: Union[str, Path], **kwargs) -> "PDFDocument":
        """
        保存文档
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            self: 支持链式调用
        """
        if not self._is_loaded:
            self.load()
        
        output_path = Path(output_path)
        
        if PYPDF2_AVAILABLE and self._pdf_reader:
            writer = PdfWriter()
            for page in self._pdf_reader.pages:
                writer.add_page(page)
            
            with open(output_path, 'wb') as f:
                writer.write(f)
        else:
            raise ValueError("没有可保存的PDF")
        
        return self
    
    def convert_to(
        self,
        target_type: DocumentType,
        output_path: Optional[Union[str, Path]] = None,
        **kwargs
    ) -> Union[str, Path, bytes]:
        """
        转换文档格式
        
        Args:
            target_type: 目标文档类型
            output_path: 输出路径（可选）
            
        Returns:
            转换后的文件路径或字节数据
        """
        if not self._is_loaded:
            self.load()
        
        if target_type == DocumentType.TXT:
            text = self.extract_text(**kwargs)
            if output_path:
                Path(output_path).write_text(text, encoding='utf-8')
                return output_path
            return text.encode('utf-8')
        
        elif target_type == DocumentType.DOCX:
            return self._convert_to_docx(output_path, **kwargs)
        
        elif target_type == DocumentType.HTML:
            return self._convert_to_html(output_path, **kwargs)
        
        elif target_type in [DocumentType.PNG, DocumentType.JPEG]:
            return self._convert_to_images(output_path, **kwargs)
        
        else:
            raise ValueError(f"不支持的目标格式: {target_type}")
    
    def _convert_to_docx(
        self,
        output_path: Optional[Union[str, Path]] = None,
        **kwargs
    ) -> Union[str, Path]:
        """转换为DOCX"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx 库: pip install python-docx")
        
        if not output_path:
            output_path = self._temp_manager.create_temp_file(suffix='.docx')
        else:
            output_path = Path(output_path)
        
        doc = Document()
        
        # 提取文本并添加到文档
        text_by_pages = self.extract_text_by_pages()
        for page_num, text in sorted(text_by_pages.items()):
            if text.strip():
                doc.add_paragraph(f"=== Page {page_num} ===")
                for line in text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                doc.add_paragraph()  # 空行
        
        doc.save(output_path)
        return output_path
    
    def _convert_to_html(
        self,
        output_path: Optional[Union[str, Path]] = None,
        **kwargs
    ) -> Union[str, Path]:
        """转换为HTML"""
        if not output_path:
            output_path = self._temp_manager.create_temp_file(suffix='.html')
        else:
            output_path = Path(output_path)
        
        html_parts = [
            '<!DOCTYPE html>',
            '<html>',
            '<head>',
            '<meta charset="UTF-8">',
            f'<title>{self._metadata.file_name or "PDF Document"}</title>',
            '<style>',
            'body { font-family: Arial, sans-serif; margin: 40px; }',
            '.page { border: 1px solid #ccc; margin: 20px 0; padding: 20px; }',
            '.page-header { font-weight: bold; margin-bottom: 10px; }',
            '</style>',
            '</head>',
            '<body>'
        ]
        
        text_by_pages = self.extract_text_by_pages()
        for page_num, text in sorted(text_by_pages.items()):
            html_parts.append(f'<div class="page">')
            html_parts.append(f'<div class="page-header">Page {page_num}</div>')
            html_parts.append('<pre>')
            html_parts.append(text.replace('<', '&lt;').replace('>', '&gt;'))
            html_parts.append('</pre>')
            html_parts.append('</div>')
        
        html_parts.extend(['</body>', '</html>'])
        
        html_content = '\n'.join(html_parts)
        output_path.write_text(html_content, encoding='utf-8')
        
        return output_path
    
    def _convert_to_images(
        self,
        output_path: Optional[Union[str, Path]] = None,
        **kwargs
    ) -> List[Path]:
        """转换为图片"""
        if not PDF2IMAGE_AVAILABLE:
            raise ImportError("需要安装 pdf2image 库: pip install pdf2image")
        
        if not output_path:
            output_dir = self._temp_manager.create_temp_dir()
        else:
            output_dir = Path(output_path)
            output_dir.mkdir(parents=True, exist_ok=True)
        
        dpi = kwargs.get('dpi', 200)
        fmt = kwargs.get('format', 'png')
        
        if self._file_path:
            images = convert_from_path(str(self._file_path), dpi=dpi, fmt=fmt)
        elif self._file_stream:
            self._file_stream.seek(0)
            images = convert_from_bytes(self._file_stream.read(), dpi=dpi, fmt=fmt)
        else:
            raise ValueError("没有可用的文件源")
        
        image_paths = []
        for idx, image in enumerate(images):
            image_path = output_dir / f"page_{idx + 1}.{fmt}"
            image.save(image_path)
            image_paths.append(image_path)
        
        return image_paths
    
    def merge_pdfs(
        self,
        other_pdfs: List[Union[str, Path, "PDFDocument"]],
        output_path: Optional[Union[str, Path]] = None
    ) -> Union[str, Path]:
        """
        合并多个PDF
        
        Args:
            other_pdfs: 其他PDF文件路径或文档对象
            output_path: 输出路径
            
        Returns:
            合并后的PDF路径
        """
        if not PYPDF2_AVAILABLE:
            raise ImportError("需要安装 PyPDF2 库")
        
        if not output_path:
            output_path = self._temp_manager.create_temp_file(suffix='.pdf')
        else:
            output_path = Path(output_path)
        
        writer = PdfWriter()
        
        # 添加当前PDF
        if self._pdf_reader:
            for page in self._pdf_reader.pages:
                writer.add_page(page)
        
        # 添加其他PDF
        for pdf in other_pdfs:
            if isinstance(pdf, PDFDocument):
                if pdf._pdf_reader:
                    for page in pdf._pdf_reader.pages:
                        writer.add_page(page)
            else:
                other_reader = PdfReader(str(pdf))
                for page in other_reader.pages:
                    writer.add_page(page)
        
        with open(output_path, 'wb') as f:
            writer.write(f)
        
        return output_path
    
    def split_pdf(
        self,
        page_ranges: List[Tuple[int, int]],
        output_dir: Optional[Union[str, Path]] = None
    ) -> List[Path]:
        """
        拆分PDF
        
        Args:
            page_ranges: 页面范围列表，每个元素为(start, end)元组（从1开始）
            output_dir: 输出目录
            
        Returns:
            拆分后的PDF路径列表
        """
        if not PYPDF2_AVAILABLE:
            raise ImportError("需要安装 PyPDF2 库")
        
        if not output_dir:
            output_dir = self._temp_manager.create_temp_dir()
        else:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
        
        output_paths = []
        
        for idx, (start, end) in enumerate(page_ranges):
            writer = PdfWriter()
            
            for page_idx in range(start - 1, min(end, self._page_count)):
                writer.add_page(self._pdf_reader.pages[page_idx])
            
            output_path = output_dir / f"split_{idx + 1}_pages_{start}-{end}.pdf"
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            output_paths.append(output_path)
        
        return output_paths
    
    def close(self) -> None:
        """关闭文档，释放资源"""
        if self._pdfplumber_doc:
            try:
                self._pdfplumber_doc.close()
            except Exception as e:
                logger.warning(f"关闭pdfplumber文档失败: {e}")
            self._pdfplumber_doc = None
        
        self._pdf_reader = None
        self._temp_manager.cleanup()
        super().close()


class PDFHandler(BaseDocumentHandler):
    """
    PDF文档处理器
    
    处理 .pdf 格式的PDF文档。
    """
    
    @property
    def supported_types(self) -> List[DocumentType]:
        """返回支持的文档类型列表"""
        return [DocumentType.PDF]
    
    @property
    def supported_extensions(self) -> List[str]:
        """返回支持的文件扩展名列表"""
        return ['.pdf']
    
    def create_document(
        self,
        file_path: Optional[Union[str, Path]] = None,
        file_stream: Optional[BinaryIO] = None
    ) -> PDFDocument:
        """
        创建PDF文档对象
        
        Args:
            file_path: 文件路径
            file_stream: 文件流
            
        Returns:
            PDF文档对象
        """
        if file_path:
            self.validate_file(file_path)
        
        return PDFDocument(file_path=file_path, file_stream=file_stream)
    
    def create_new_document(self) -> PDFDocument:
        """
        创建新的空白PDF文档
        
        Returns:
            新的PDF文档对象
        """
        if not PYPDF2_AVAILABLE:
            raise ImportError("需要安装 PyPDF2 库")
        
        doc = PDFDocument()
        # 创建一个空白的PDF
        from PyPDF2 import PdfWriter
        writer = PdfWriter()
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        doc._pdf_reader = PdfReader(buffer)
        doc._page_count = 0
        doc._document_type = DocumentType.PDF
        doc._is_loaded = True
        return doc


# 便捷函数
def extract_text_from_pdf(file_path: Union[str, Path], **kwargs) -> str:
    """
    从PDF文档提取文本
    
    Args:
        file_path: PDF文档路径
        **kwargs: 其他参数
        
    Returns:
        文档文本内容
    """
    handler = PDFHandler()
    doc = handler.create_document(file_path=file_path)
    with doc:
        return doc.extract_text(**kwargs)


def extract_tables_from_pdf(file_path: Union[str, Path], **kwargs) -> List[List[List[str]]]:
    """
    从PDF文档提取表格
    
    Args:
        file_path: PDF文档路径
        **kwargs: 其他参数
        
    Returns:
        表格数据列表
    """
    handler = PDFHandler()
    doc = handler.create_document(file_path=file_path)
    with doc:
        return doc.extract_tables(**kwargs)


def merge_pdf_files(
    pdf_files: List[Union[str, Path]],
    output_path: Union[str, Path]
) -> Path:
    """
    合并多个PDF文件
    
    Args:
        pdf_files: PDF文件路径列表
        output_path: 输出路径
        
    Returns:
        合并后的PDF路径
    """
    if not PYPDF2_AVAILABLE:
        raise ImportError("需要安装 PyPDF2 库")
    
    writer = PdfWriter()
    
    for pdf_file in pdf_files:
        reader = PdfReader(str(pdf_file))
        for page in reader.pages:
            writer.add_page(page)
    
    output_path = Path(output_path)
    with open(output_path, 'wb') as f:
        writer.write(f)
    
    return output_path
