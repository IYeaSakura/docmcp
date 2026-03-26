"""
Word文档处理器模块

支持 .doc 和 .docx 格式的Word文档处理。
"""

import io
import logging
import os
import re
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, BinaryIO, Dict, Iterator, List, Optional, Tuple, Union
from urllib.parse import unquote

from .base import BaseDocument, BaseDocumentHandler, BatchProcessor
from ..document import DocumentMetadata, DocumentType, ExtractedContent
from ..utils import TempFileManager, detect_file_type, temp_file_context

# 配置日志
logger = logging.getLogger(__name__)

# 尝试导入依赖库
try:
    import docx
    from docx import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 未安装，DOCX功能受限")

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False
    logger.warning("docx2txt 未安装")

try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False
    logger.warning("olefile 未安装，旧版DOC功能受限")


class WordDocument(BaseDocument):
    """
    Word文档对象
    
    支持 .doc 和 .docx 格式的Word文档。
    """
    
    def __init__(
        self,
        file_path: Optional[Union[str, Path]] = None,
        file_stream: Optional[BinaryIO] = None,
        document_type: Optional[DocumentType] = None
    ):
        """
        初始化Word文档对象
        
        Args:
            file_path: 文档文件路径
            file_stream: 文档文件流
            document_type: 文档类型
        """
        super().__init__(file_path, file_stream, document_type)
        self._docx_doc: Optional[Any] = None
        self._temp_manager = TempFileManager()
    
    def load(self, **kwargs) -> "WordDocument":
        """
        加载Word文档
        
        Returns:
            self: 支持链式调用
            
        Raises:
            FileNotFoundError: 文件不存在
            PermissionError: 没有读取权限
            ValueError: 文件格式错误
        """
        if self._is_loaded:
            return self
        
        try:
            # 确定文档类型
            if self._document_type == DocumentType.UNKNOWN:
                if self._file_path:
                    self._document_type = detect_file_type(self._file_path)
                elif self._file_stream:
                    self._document_type = detect_file_type(self._file_stream)
            
            # 加载文档
            if self._document_type == DocumentType.DOCX:
                self._load_docx()
            elif self._document_type == DocumentType.DOC:
                self._load_doc()
            else:
                # 尝试自动检测
                self._try_auto_load()
            
            self._is_loaded = True
            self._update_metadata_from_file()
            
        except Exception as e:
            logger.error(f"加载Word文档失败: {e}")
            raise ValueError(f"无法加载Word文档: {e}")
        
        return self
    
    def _load_docx(self) -> None:
        """加载DOCX格式文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx 库: pip install python-docx")
        
        if self._file_stream:
            self._file_stream.seek(0)
            self._docx_doc = DocxDocument(self._file_stream)
        elif self._file_path:
            self._docx_doc = DocxDocument(self._file_path)
        else:
            raise ValueError("没有可用的文件源")
    
    def _load_doc(self) -> None:
        """加载DOC格式文档（旧版）"""
        # 旧版DOC需要转换为DOCX或使用antiword等工具
        # 这里尝试使用textract或转换为DOCX
        try:
            docx_path = self._convert_doc_to_docx()
            self._docx_doc = DocxDocument(docx_path)
            self._document_type = DocumentType.DOCX  # 转换后类型改变
        except Exception as e:
            logger.error(f"加载DOC文件失败: {e}")
            raise
    
    def _try_auto_load(self) -> None:
        """尝试自动检测并加载"""
        errors = []
        
        # 尝试DOCX
        try:
            self._load_docx()
            self._document_type = DocumentType.DOCX
            return
        except Exception as e:
            errors.append(f"DOCX: {e}")
        
        # 尝试DOC
        try:
            self._load_doc()
            self._document_type = DocumentType.DOC
            return
        except Exception as e:
            errors.append(f"DOC: {e}")
        
        raise ValueError(f"无法识别文档格式: {'; '.join(errors)}")
    
    def _convert_doc_to_docx(self) -> Path:
        """
        将DOC转换为DOCX
        
        Returns:
            转换后的DOCX文件路径
        """
        if not self._file_path:
            raise ValueError("需要文件路径才能转换DOC")
        
        # 尝试使用LibreOffice转换
        try:
            return self._convert_with_libreoffice()
        except Exception as e:
            logger.warning(f"LibreOffice转换失败: {e}")
        
        # 尝试使用antiword提取文本
        try:
            return self._convert_with_antiword()
        except Exception as e:
            logger.warning(f"Antiword转换失败: {e}")
        
        raise ValueError("无法转换DOC文件，请安装LibreOffice或antiword")
    
    def _convert_with_libreoffice(self) -> Path:
        """使用LibreOffice转换"""
        if not self._file_path:
            raise ValueError("需要文件路径")
        
        output_dir = self._temp_manager.create_temp_dir()
        
        cmd = [
            'soffice',
            '--headless',
            '--convert-to', 'docx',
            '--outdir', str(output_dir),
            str(self._file_path)
        ]
        
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60
        )
        
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice转换失败: {result.stderr}")
        
        # 找到转换后的文件
        output_file = output_dir / f"{self._file_path.stem}.docx"
        if output_file.exists():
            return output_file
        
        # 尝试其他文件名
        for f in output_dir.glob("*.docx"):
            return f
        
        raise FileNotFoundError("转换后的文件未找到")
    
    def _convert_with_antiword(self) -> Path:
        """使用antiword提取文本并创建简单DOCX"""
        if not self._file_path:
            raise ValueError("需要文件路径")
        
        cmd = ['antiword', str(self._file_path)]
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode != 0:
            raise RuntimeError(f"Antiword提取失败: {result.stderr}")
        
        # 创建简单DOCX
        output_path = self._temp_manager.create_temp_file(suffix='.docx')
        doc = DocxDocument()
        
        for line in result.stdout.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(output_path)
        return output_path
    
    def extract_text(self, **kwargs) -> str:
        """
        提取纯文本内容
        
        Args:
            preserve_whitespace: 是否保留空白字符
            
        Returns:
            文档的纯文本内容
        """
        if not self._is_loaded:
            self.load()
        
        preserve_whitespace = kwargs.get('preserve_whitespace', False)
        
        if self._docx_doc:
            texts = []
            for para in self._docx_doc.paragraphs:
                text = para.text
                if text.strip() or preserve_whitespace:
                    texts.append(text)
            
            # 提取表格中的文本
            for table in self._docx_doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        texts.append(' | '.join(row_texts))
            
            return '\n'.join(texts)
        
        return ""
    
    def extract_content(self, **kwargs) -> ExtractedContent:
        """
        提取完整内容
        
        Args:
            include_images: 是否包含图片
            include_hyperlinks: 是否包含超链接
            
        Returns:
            包含文本、表格、图片等的完整内容
        """
        if not self._is_loaded:
            self.load()
        
        include_images = kwargs.get('include_images', False)
        include_hyperlinks = kwargs.get('include_hyperlinks', True)
        
        content = ExtractedContent()
        
        if not self._docx_doc:
            return content
        
        # 提取段落
        for para in self._docx_doc.paragraphs:
            if para.text.strip():
                content.paragraphs.append(para.text)
        
        content.text = '\n'.join(content.paragraphs)
        
        # 提取表格
        for table in self._docx_doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            content.tables.append(table_data)
        
        # 提取标题
        for para in self._docx_doc.paragraphs:
            if para.style and para.style.name and 'Heading' in para.style.name:
                level = 0
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                except ValueError:
                    pass
                
                content.headings.append({
                    'text': para.text,
                    'level': level,
                    'style': para.style.name
                })
        
        # 提取超链接
        if include_hyperlinks:
            content.hyperlinks = self._extract_hyperlinks()
        
        # 提取图片
        if include_images:
            content.images = self._extract_images()
        
        return content
    
    def _extract_hyperlinks(self) -> List[Dict[str, str]]:
        """提取超链接"""
        hyperlinks = []
        
        if not self._docx_doc:
            return hyperlinks
        
        # 从段落中提取超链接
        for para in self._docx_doc.paragraphs:
            for run in para.runs:
                # 检查是否有超链接
                if run._element.xpath('.//a:href', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                    href = run._element.xpath('.//a:href', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                    if href:
                        hyperlinks.append({
                            'text': run.text,
                            'url': href[0]
                        })
        
        # 从关系部分提取超链接
        try:
            rels = self._docx_doc.part.rels
            for rel in rels.values():
                if "hyperlink" in rel.target_ref:
                    hyperlinks.append({
                        'text': rel.target_ref,
                        'url': rel.target_ref
                    })
        except Exception as e:
            logger.debug(f"提取超链接失败: {e}")
        
        return hyperlinks
    
    def _extract_images(self) -> List[Dict[str, Any]]:
        """提取图片"""
        images = []
        
        if not self._docx_doc:
            return images
        
        try:
            # 获取文档中的所有图片
            image_parts = []
            for rel in self._docx_doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_parts.append(rel.target_part)
            
            for i, image_part in enumerate(image_parts):
                image_data = {
                    'index': i,
                    'content_type': getattr(image_part, 'content_type', 'unknown'),
                    'size': len(image_part.blob) if hasattr(image_part, 'blob') else 0,
                }
                
                # 保存图片数据
                if hasattr(image_part, 'blob'):
                    image_data['data'] = image_part.blob
                
                images.append(image_data)
                
        except Exception as e:
            logger.warning(f"提取图片失败: {e}")
        
        return images
    
    def extract_metadata(self, **kwargs) -> DocumentMetadata:
        """
        提取文档元数据
        
        Returns:
            文档元数据对象
        """
        if not self._is_loaded:
            self.load()
        
        metadata = DocumentMetadata()
        
        # 从文件系统获取基本元数据
        self._update_metadata_from_file()
        metadata.file_size = self._metadata.file_size
        metadata.file_path = self._metadata.file_path
        metadata.file_name = self._metadata.file_name
        metadata.file_extension = self._metadata.file_extension
        
        if not self._docx_doc:
            return metadata
        
        # 提取Core属性
        try:
            core_props = self._docx_doc.core_properties
            metadata.title = core_props.title
            metadata.author = core_props.author
            metadata.subject = core_props.subject
            metadata.keywords = core_props.keywords
            metadata.comments = core_props.comments
            metadata.last_modified_by = core_props.last_modified_by
            metadata.revision = core_props.revision
            metadata.category = core_props.category
            
            # 日期处理
            if core_props.created:
                metadata.created = core_props.created
            if core_props.modified:
                metadata.modified = core_props.modified
                
        except Exception as e:
            logger.debug(f"提取元数据失败: {e}")
        
        return metadata
    
    def save(self, output_path: Union[str, Path], **kwargs) -> "WordDocument":
        """
        保存文档
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            self: 支持链式调用
        """
        if not self._is_loaded:
            self.load()
        
        output_path = Path(output_path)
        
        if self._docx_doc:
            self._docx_doc.save(output_path)
        else:
            raise ValueError("没有可保存的文档")
        
        return self
    
    def convert_to(
        self,
        target_type: DocumentType,
        output_path: Optional[Union[str, Path]] = None,
        **kwargs
    ) -> Union[str, Path, bytes]:
        """
        转换文档格式
        
        Args:
            target_type: 目标文档类型
            output_path: 输出路径（可选）
            
        Returns:
            转换后的文件路径或字节数据
        """
        if not self._is_loaded:
            self.load()
        
        if target_type == DocumentType.TXT:
            text = self.extract_text()
            if output_path:
                Path(output_path).write_text(text, encoding='utf-8')
                return output_path
            return text.encode('utf-8')
        
        elif target_type == DocumentType.PDF:
            return self._convert_to_pdf(output_path, **kwargs)
        
        elif target_type in [DocumentType.DOCX, DocumentType.DOC]:
            if output_path:
                self.save(output_path)
                return output_path
            else:
                # 返回字节数据
                buffer = io.BytesIO()
                self._docx_doc.save(buffer)
                return buffer.getvalue()
        
        else:
            raise ValueError(f"不支持的目标格式: {target_type}")
    
    def _convert_to_pdf(
        self,
        output_path: Optional[Union[str, Path]] = None,
        **kwargs
    ) -> Union[str, Path]:
        """
        转换为PDF
        
        Args:
            output_path: 输出路径
            
        Returns:
            PDF文件路径
        """
        if not output_path:
            output_path = self._temp_manager.create_temp_file(suffix='.pdf')
        else:
            output_path = Path(output_path)
        
        # 使用LibreOffice转换
        try:
            # 先保存为临时DOCX
            temp_docx = self._temp_manager.create_temp_file(suffix='.docx')
            self._docx_doc.save(temp_docx)
            
            cmd = [
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_path.parent),
                str(temp_docx)
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"PDF转换失败: {result.stderr}")
            
            # 重命名输出文件
            expected_output = output_path.parent / f"{temp_docx.stem}.pdf"
            if expected_output.exists() and expected_output != output_path:
                expected_output.rename(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"PDF转换失败: {e}")
            raise
    
    def close(self) -> None:
        """关闭文档，释放资源"""
        self._docx_doc = None
        self._temp_manager.cleanup()
        super().close()


class WordHandler(BaseDocumentHandler):
    """
    Word文档处理器
    
    处理 .doc 和 .docx 格式的Word文档。
    """
    
    @property
    def supported_types(self) -> List[DocumentType]:
        """返回支持的文档类型列表"""
        return [DocumentType.DOC, DocumentType.DOCX]
    
    @property
    def supported_extensions(self) -> List[str]:
        """返回支持的文件扩展名列表"""
        return ['.doc', '.docx']
    
    def create_document(
        self,
        file_path: Optional[Union[str, Path]] = None,
        file_stream: Optional[BinaryIO] = None
    ) -> WordDocument:
        """
        创建Word文档对象
        
        Args:
            file_path: 文件路径
            file_stream: 文件流
            
        Returns:
            Word文档对象
        """
        if file_path:
            self.validate_file(file_path)
        
        return WordDocument(file_path=file_path, file_stream=file_stream)
    
    def create_new_document(self) -> WordDocument:
        """
        创建新的空白Word文档
        
        Returns:
            新的Word文档对象
        """
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx 库")
        
        doc = WordDocument()
        doc._docx_doc = DocxDocument()
        doc._document_type = DocumentType.DOCX
        doc._is_loaded = True
        return doc


# 便捷函数
def extract_text_from_word(file_path: Union[str, Path]) -> str:
    """
    从Word文档提取文本
    
    Args:
        file_path: Word文档路径
        
    Returns:
        文档文本内容
    """
    handler = WordHandler()
    doc = handler.create_document(file_path=file_path)
    with doc:
        return doc.extract_text()


def extract_tables_from_word(file_path: Union[str, Path]) -> List[List[List[str]]]:
    """
    从Word文档提取表格
    
    Args:
        file_path: Word文档路径
        
    Returns:
        表格数据列表
    """
    handler = WordHandler()
    doc = handler.create_document(file_path=file_path)
    with doc:
        content = doc.extract_content()
        return content.tables
